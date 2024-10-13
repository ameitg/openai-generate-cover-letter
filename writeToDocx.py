from docx import Document

def writeToDocx(filename, content):
    # Send a GET request to the webpage


    # Create a new Document
    doc = Document()

    # Add a title
    # doc.add_heading('Document Title', 0)

    # Add a paragraph
    doc.add_paragraph(content)

    # Save the document
    doc.save(filename)

